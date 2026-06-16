# -*- coding: utf-8 -*-
"""Sinh file Word: hướng dẫn lấy DB từ VPS về máy dev."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC = Document()

# ── base styles ──
normal = DOC.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

def shade(paragraph, fill='F2F2F2'):
    pPr = paragraph._p.get_or_add_pPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), fill)
    pPr.append(sh)

def code(text):
    """Khối lệnh: nền xám, font Consolas."""
    p = DOC.add_paragraph()
    p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    shade(p)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    return p

def note(text):
    p = DOC.add_paragraph()
    run = p.add_run('Lưu ý: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xB0, 0x45, 0x00)
    p.add_run(text)
    return p

def para(text, bold=False):
    p = DOC.add_paragraph()
    r = p.add_run(text); r.bold = bold
    return p

def bullet(text):
    DOC.add_paragraph(text, style='List Bullet')

# ── Title ──
t = DOC.add_heading('HƯỚNG DẪN LẤY DATABASE TỪ VPS VỀ MÁY DEV', level=0)
DOC.add_paragraph('Dự án: Hello world (TECAPRO) — PostgreSQL. Cập nhật: 15/06/2026.')

# ── 0. Tổng quan ──
DOC.add_heading('0. Ý tưởng & yêu cầu', level=1)
para('Quy trình gồm 3 việc: (1) tạo bản dump DB trên VPS, (2) tải file dump về máy dev, '
     '(3) phục hồi (restore) vào DB local "hello_web_db". File dump KHÔNG chứa file đính kèm '
     '(thư mục server/uploads) — nếu cần cả file thì copy riêng (xem Bước 6).')
para('Cần có sẵn trên máy dev:', bold=True)
bullet('PostgreSQL client (pg_dump, pg_restore, psql) — máy đang có PostgreSQL 18 trong PATH.')
bullet('Công cụ SSH/SCP: OpenSSH (ssh, scp) hoặc PuTTY (plink, pscp) — máy đang có cả hai.')
para('Thông tin cần điền (lấy từ VPS — xem Bước 1):', bold=True)
bullet('VPS_HOST: địa chỉ IP/tên miền của VPS.')
bullet('SSH_USER: tài khoản SSH (vd root hoặc deploy).')
bullet('VPS_DB / VPS_DB_USER / VPS_DB_PASS: tên DB, user, mật khẩu Postgres trên VPS.')

# ── 1. Lấy thông tin kết nối trên VPS ──
DOC.add_heading('1. Lấy thông tin kết nối DB trên VPS', level=1)
para('Đăng nhập VPS qua SSH (mở Git Bash hoặc PowerShell trên máy dev):')
code('ssh SSH_USER@VPS_HOST')
para('Trong VPS, mở file .env của ứng dụng để xem chuỗi kết nối (đường dẫn tùy nơi đặt app, '
     'ví dụ /var/www/hello-world hoặc ~/app):')
code('grep DATABASE_URL /duong/dan/toi/app/.env')
para('Kết quả có dạng:')
code('DATABASE_URL=postgres://VPS_DB_USER:VPS_DB_PASS@localhost:5432/VPS_DB')
note('Ghi lại VPS_DB_USER, VPS_DB_PASS, VPS_DB từ dòng này. DB chạy ngay trên VPS nên host là '
     'localhost (đứng từ trong VPS).')

# ── 2. Dump trên VPS ──
DOC.add_heading('2. Tạo bản dump trên VPS', level=1)
para('Vẫn trong phiên SSH ở VPS, tạo dump định dạng custom (nén, phục hồi linh hoạt):')
code('pg_dump "postgres://VPS_DB_USER:VPS_DB_PASS@localhost:5432/VPS_DB" -Fc -f /tmp/hello_web_db.dump')
para('Kiểm tra file đã tạo:')
code('ls -lh /tmp/hello_web_db.dump')
note('-Fc = custom format (khuyến nghị, dùng với pg_restore). Nếu muốn file SQL thuần để đọc/sửa '
     'tay thì bỏ -Fc và đổi đuôi .sql, sau này restore bằng psql -f thay vì pg_restore.')

# ── 3. Tải file về máy dev ──
DOC.add_heading('3. Tải file dump về máy dev', level=1)
para('MỞ CỬA SỔ MỚI trên máy dev (thoát SSH hoặc dùng terminal khác). Chọn 1 trong 2 cách:')
para('Cách A — OpenSSH (scp):', bold=True)
code('scp SSH_USER@VPS_HOST:/tmp/hello_web_db.dump "D:/backup/hello_web_db.dump"')
para('Cách B — PuTTY (pscp):', bold=True)
code('pscp SSH_USER@VPS_HOST:/tmp/hello_web_db.dump "D:\\backup\\hello_web_db.dump"')
note('Tạo trước thư mục D:\\backup (hoặc nơi bạn muốn) để chứa file. File có thể vài chục MB.')

# ── 4. Sao lưu DB local trước khi ghi đè ──
DOC.add_heading('4. Sao lưu DB local hiện tại (nên làm)', level=1)
para('Bước phục hồi sẽ XÓA và ghi đè DB local. Hãy backup trước cho an toàn:')
code('pg_dump "postgres://postgres:MAT_KHAU_LOCAL@localhost:5432/hello_web_db" -Fc '
     '-f "D:/backup/local_truoc_khi_restore.dump"')
note('Thay MAT_KHAU_LOCAL bằng mật khẩu Postgres trên máy dev (xem .env local: '
     'DATABASE_URL=postgres://postgres:...@localhost:5432/hello_web_db).')

# ── 5. Phục hồi vào DB local ──
DOC.add_heading('5. Phục hồi (restore) vào DB local', level=1)
para('5.1 — TẮT backend dev đang chạy (Ctrl+C ở cửa sổ npm run dev). DB không được có kết nối '
     'đang mở thì mới xóa/tạo lại được.', bold=True)
para('5.2 — Xóa và tạo lại DB local rỗng (để bản copy sạch, không lẫn dữ liệu cũ):')
code('psql -U postgres -h localhost -c "DROP DATABASE IF EXISTS hello_web_db;"\n'
     'psql -U postgres -h localhost -c "CREATE DATABASE hello_web_db OWNER postgres;"')
para('5.3 — Nạp dữ liệu từ file dump của VPS:')
code('pg_restore -U postgres -h localhost -d hello_web_db --no-owner --no-privileges '
     '"D:/backup/hello_web_db.dump"')
note('--no-owner --no-privileges để bỏ qua chủ sở hữu/role của VPS (máy dev dùng user postgres), '
     'tránh lỗi "role does not exist". Nếu pg_restore hỏi mật khẩu, nhập mật khẩu postgres local; '
     'hoặc đặt sẵn: trên PowerShell  $env:PGPASSWORD="..."  / trên Git Bash  export PGPASSWORD=...')
para('Nếu dump là file .sql thuần (không -Fc) thì thay 5.3 bằng:')
code('psql -U postgres -h localhost -d hello_web_db -f "D:/backup/hello_web_db.sql"')

# ── 6. Lấy kèm file đính kèm (uploads) ──
DOC.add_heading('6. (Tùy chọn) Lấy kèm file đính kèm', level=1)
para('File tải lên (đính kèm công việc, tài liệu) nằm ở thư mục server/uploads trên VPS, '
     'KHÔNG nằm trong DB và KHÔNG có trong git. Nếu cần đầy đủ để chạy thử, copy về:')
code('scp -r SSH_USER@VPS_HOST:/duong/dan/toi/app/server/uploads/* '
     '"D:/Cursor Learning/Hello world/server/uploads/"')
note('Thư mục server/uploads bị git bỏ qua — đây là bản sao để chạy local, đừng commit lên git.')

# ── 7. Kiểm tra ──
DOC.add_heading('7. Kiểm tra sau khi phục hồi', level=1)
para('Đếm bảng và thử vài bảng chính:')
code('psql -U postgres -h localhost -d hello_web_db -c "\\dt"\n'
     'psql -U postgres -h localhost -d hello_web_db -c "SELECT count(*) FROM app_user;"\n'
     'psql -U postgres -h localhost -d hello_web_db -c "SELECT count(*) FROM contract_out;"')
para('Bật lại backend và đăng nhập thử bằng tài khoản có trên VPS:')
code('npm run dev')

# ── 8. Cách thay thế: SSH tunnel ──
DOC.add_heading('8. Cách thay thế — dump trực tiếp qua SSH tunnel', level=1)
para('Nếu không muốn tạo file trên VPS, có thể mở "đường hầm" rồi dump thẳng từ máy dev.')
para('8.1 — Mở tunnel (giữ cửa sổ này chạy, cổng local 5433 trỏ tới Postgres trên VPS):')
code('ssh -L 5433:localhost:5432 SSH_USER@VPS_HOST')
para('8.2 — Cửa sổ KHÁC trên máy dev, dump qua cổng 5433:')
code('pg_dump "postgres://VPS_DB_USER:VPS_DB_PASS@localhost:5433/VPS_DB" -Fc '
     '-f "D:/backup/hello_web_db.dump"')
para('Sau đó phục hồi như Bước 5.')

# ── 9. Lỗi thường gặp ──
DOC.add_heading('9. Lỗi thường gặp', level=1)
bullet('"database is being accessed by other users" khi DROP: còn kết nối mở — tắt backend dev, '
       'đóng các tab psql/DBeaver đang nối tới hello_web_db, rồi thử lại.')
bullet('"role ... does not exist" khi restore: đã thêm --no-owner --no-privileges chưa? '
       'Đây chính là cách xử lý.')
bullet('"password authentication failed": sai mật khẩu — kiểm tra lại .env (local hoặc VPS) '
       'hoặc dùng biến PGPASSWORD.')
bullet('Phiên bản pg_dump khác server: nên dump bằng pg_dump CHẠY TRÊN VPS (Bước 2) để khớp '
       'phiên bản server ở đó; restore vào PostgreSQL 18 local thì luôn được.')
bullet('"version mismatch" khi dump qua tunnel: pg_dump 18 trên máy dev vẫn dump được server cũ hơn; '
       'nếu vẫn lỗi, dùng cách Bước 2 (dump ngay trên VPS).')

OUT = r'd:/Cursor Learning/Hello world/docs/Huong-dan-lay-DB-tu-VPS.docx'
DOC.save(OUT)
print('Saved:', OUT)
